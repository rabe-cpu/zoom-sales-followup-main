#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


YELLOW = "FFF2CC"
BLUE = "D9EAF7"
GREEN = "E2F0D9"
INLINE_YELLOW_WITH_LABEL = re.compile(r"\[黄色:([^\]]*)\]([^\[]*?)\[/黄色\]")
INLINE_YELLOW_NO_LABEL_CLOSE = re.compile(r"\[黄色:([^\[]*?)\[/黄色\]")
BLOCK_YELLOW_START = re.compile(r"^\[黄色:[^\]]+\]$")

KEY_LABELS = {
    "overallSummary": "総合所見",
    "meetingOutcome": "今の着地点",
    "benchmarkGapAnalysis": "トップ営業との差分",
    "keyDifference": "決定的な差分",
    "missingActions": "次に補う動き",
    "phase": "場面",
    "missing": "不足していた動き",
    "impact": "起きやすい影響",
    "hiddenNeeds": "表に出ていない本音",
    "unspokenIssues": "言い切れていない不安",
    "competingServices": "比較対象",
    "decisionCriteria": "判断基準",
    "buyingTemperature": "温度感",
    "name": "テーマ",
    "description": "見立て",
    "counterMeasure": "次の打ち手",
    "topic": "論点",
    "gap": "ズレ",
    "solution": "整え方",
    "item": "観点",
    "detail": "指導メモ",
    "theme": "テーマ",
    "scene": "場面",
    "insight": "トップ営業の見立て",
    "issue": "詰まりやすい点",
    "strategy": "次の打ち手",
    "script": "そのまま使える台詞",
    "outcome": "狙う着地",
    "useCase": "使う場面",
    "customerSignals": "顧客シグナル",
    "topSalesMove": "トップ営業の動き",
    "copyTalk": "そのまま使える台詞",
    "whyItWorks": "効く理由",
    "ngExamples": "避ける言い方",
    "practice": "次回実践",
    "currentPhase": "現在フェーズ",
    "currentGoal": "この接点の目的",
    "keepUntilLater": "今は言い切らないこと",
    "mustHearBeforeProposal": "次に必ず聞くこと",
    "planDecisionPath": "判断材料の並べ方",
    "nextBestAction": "次の一手",
    "hearingQuestions": "次に聞く質問",
    "recommendedAnswer": "返信が来た時の返答案",
    "benchmarkPattern": "使った型",
    "delivery": "伝え方",
    "sourceMoment": "根拠になった商談場面",
    "recommendedTalk": "接続トーク",
    "effectiveQuestions": "効く質問",
    "effectiveReplies": "効く返答",
    "avoidedTalk": "避ける言い方",
}

FEEDBACK_TITLES = {
    "総合概要",
    "顧客インサイト",
    "認知バイアス",
    "期待値のズレ",
    "良かった点",
    "改善ポイント",
    "AIコーチングカード",
    "再現する勝ち筋",
    "商談フェーズ",
    "顧客シグナル",
    "次の一手",
    "ベンチマーク営業再現",
    "文脈接続",
    "属性別対応",
}


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Yu Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    style.font.size = Pt(10.5)


def add_table(document: Document, headers: list[str], rows: list[list[str]], fill: str) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        shade_cell(cell, fill)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)


def stringify(value: Any) -> str:
    if isinstance(value, list):
        return "\n".join(str(item) for item in value)
    if isinstance(value, dict):
        return "\n".join(f"{key}: {item}" for key, item in value.items())
    return "" if value is None else str(value)


def label_for(key: Any) -> str:
    return KEY_LABELS.get(str(key), str(key))


def compact_sentence(value: Any) -> str:
    if isinstance(value, list):
        return "、".join(compact_sentence(item) for item in value if item not in (None, "", [], {}))
    if isinstance(value, dict):
        parts = []
        for key, item in value.items():
            if item in (None, "", [], {}):
                continue
            parts.append(f"{label_for(key)}は{compact_sentence(item)}")
        return "。".join(parts)
    return str(value).strip()


def add_labeled_paragraph(document: Document, label: str, text: Any) -> None:
    if text in (None, "", [], {}):
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(compact_sentence(text))


def item_title(item: dict[str, Any], fallback: str) -> str:
    for key in ("title", "name", "theme", "item", "topic", "勝ち筋名"):
        if item.get(key):
            return str(item[key])
    return fallback


def add_coaching_intro(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("トップ営業の指導: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    paragraph.add_run(text)


def add_feedback_dict(document: Document, value: dict[str, Any], level: int = 2) -> None:
    for index, (key, item) in enumerate(value.items(), start=1):
        if item in (None, "", [], {}):
            continue
        document.add_heading(label_for(key), level=min(level, 4))
        if isinstance(item, dict):
            prose = compact_sentence(item)
            if prose:
                add_coaching_intro(document, f"ここは項目を読むだけではなく、営業担当が次にどう動くかまで落とし込みます。{prose}。")
        elif isinstance(item, list):
            for list_index, list_item in enumerate(item, start=1):
                if isinstance(list_item, dict):
                    title = item_title(list_item, f"{label_for(key)} {list_index}")
                    document.add_heading(title, level=min(level + 1, 4))
                    add_coaching_intro(document, f"この場面では、まず顧客の反応を受け止めてから判断軸へ戻します。{compact_sentence(list_item)}。")
                else:
                    add_coaching_intro(document, str(list_item))
        else:
            add_coaching_intro(document, str(item))


def add_structured_value(document: Document, value: Any, level: int = 2) -> None:
    if value is None or value == "":
        return
    if isinstance(value, dict):
        for key, item in value.items():
            if item is None or item == "" or item == [] or item == {}:
                continue
            document.add_heading(label_for(key), level=min(level, 4))
            add_structured_value(document, item, level + 1)
        return
    if isinstance(value, list):
        for item in value:
            if isinstance(item, dict):
                add_structured_value(document, item, level)
            else:
                document.add_paragraph(str(item), style="List Bullet")
        return
    add_body(document, str(value).splitlines())


def add_key_values(document: Document, title: str, values: dict[str, Any], fill: str = BLUE) -> None:
    if not values:
        return
    document.add_heading(title, level=1)
    add_table(document, ["項目", "内容"], [[str(k), stringify(v)] for k, v in values.items()], fill)


def add_text_section(document: Document, title: str, value: Any) -> None:
    if not value:
        return
    document.add_heading(title, level=1)
    if title.startswith("商談フィードバック") and isinstance(value, dict):
        add_deal_feedback(document, value)
        return
    add_structured_value(document, value, level=2)


def add_deal_feedback(document: Document, feedback: dict[str, Any]) -> None:
    document.add_paragraph(
        "このフィードバックは、評価ログではなく、営業担当が次回接点でそのまま使える指導メモとして読む。"
        "内部キーやスコアではなく、トップ営業が横で助言する口調で整理する。"
    )
    for section, value in feedback.items():
        if value in (None, "", [], {}):
            continue
        document.add_heading(label_for(section), level=2)
        if section in FEEDBACK_TITLES and isinstance(value, dict):
            if section == "総合概要":
                summary = value.get("overallSummary") or value.get("総合所見")
                outcome = value.get("meetingOutcome") or value.get("今の着地点")
                gap = value.get("benchmarkGapAnalysis") or value.get("トップ営業との差分")
                if summary:
                    add_coaching_intro(document, f"まず、この商談はこう見ます。{summary}")
                if outcome:
                    add_coaching_intro(document, f"今の着地点は、{outcome}です。ここを前提に、次の接点では無理に詰めず、判断材料を整える動きに寄せます。")
                if gap:
                    add_coaching_intro(document, f"トップ営業との差分はここです。{compact_sentence(gap)}。次回はこの不足分を先に埋めてください。")
                continue
            add_feedback_dict(document, value, level=3)
        elif isinstance(value, list):
            for index, item in enumerate(value, start=1):
                if isinstance(item, dict):
                    title = item_title(item, f"{section} {index}")
                    document.add_heading(title, level=3)
                    add_coaching_intro(document, f"ここで見るべきポイントは、{compact_sentence(item)}。次回はこの見立てをそのまま質問と返答に変えてください。")
                else:
                    add_coaching_intro(document, str(item))
        else:
            add_coaching_intro(document, str(value))


def add_highlighted_run(paragraph, text: str) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def add_paragraph_with_inline_yellow(document: Document, text: str, block_yellow: bool = False) -> None:
    paragraph = document.add_paragraph()
    if block_yellow:
        add_highlighted_run(paragraph, text)
        return

    cursor = 0
    matched = False
    for match in INLINE_YELLOW_WITH_LABEL.finditer(text):
        matched = True
        paragraph.add_run(text[cursor:match.start()])
        yellow_text = match.group(2) or match.group(1)
        add_highlighted_run(paragraph, yellow_text)
        cursor = match.end()
    remainder = text[cursor:]

    rebuilt = INLINE_YELLOW_NO_LABEL_CLOSE.search(remainder)
    if rebuilt:
        matched = True
        paragraph.add_run(remainder[:rebuilt.start()])
        add_highlighted_run(paragraph, rebuilt.group(1))
        paragraph.add_run(remainder[rebuilt.end():])
    elif matched:
        paragraph.add_run(remainder)
    else:
        paragraph.add_run(text)


def add_body(document: Document, body: list[Any]) -> None:
    block_yellow = False
    for raw in body:
        text = str(raw)
        stripped = text.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if BLOCK_YELLOW_START.match(stripped):
            block_yellow = True
            continue
        if stripped == "[/黄色]":
            block_yellow = False
            continue
        add_paragraph_with_inline_yellow(document, text, block_yellow)


def build_document(data: dict[str, Any], output: Path) -> None:
    document = Document()
    set_font(document)

    title = data.get("title") or "営業後送付メール"
    heading = document.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if data.get("subject"):
        p = document.add_paragraph()
        run = p.add_run(f"件名: {data['subject']}")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    if data.get("recipient"):
        document.add_paragraph(str(data["recipient"]))

    document.add_heading("営業後送付メール", level=1)
    add_body(document, data.get("body", []))

    yellow_fields = data.get("yellow_fields", [])
    if yellow_fields:
        document.add_heading("営業チェック欄", level=1)
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for index, header in enumerate(["区分", "黄色チェック項目", "確認結果"]):
            cell = table.rows[0].cells[index]
            cell.text = header
            shade_cell(cell, YELLOW)
        for field in yellow_fields:
            cells = table.add_row().cells
            cells[0].text = "営業確認"
            cells[1].text = str(field)
            cells[2].text = "未確認"
            shade_cell(cells[1], YELLOW)

    include_internal = data.get("include_internal_sections", True)
    if include_internal:
        add_text_section(
            document,
            "ソーシャルスタイル別全文メール案（社内確認用）",
            data.get("social_style_variants") or data.get("social_style_emails"),
        )
        add_text_section(
            document,
            "商談フィードバック要素",
            data.get("deal_feedback") or data.get("sales_feedback"),
        )

    final_check = data.get("final_check", {}) if include_internal else {}
    if final_check:
        add_key_values(document, "最終確認", final_check, GREEN)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def sample_data() -> dict[str, Any]:
    return {
        "title": "営業後送付メール サンプル",
        "subject": "本日のご面談のお礼",
        "recipient": "〇〇さま",
        "body": [
            "本日はお時間をいただき、ありがとうございました。",
            "[黄色:営業確認]",
            "お話の中で、現場での運用負荷を抑えながら進めたいという点が大事なのだと受け取りました。",
            "[/黄色]",
            "まずはお時間があるときに、下記をご確認いただけると幸いです。",
            "ZOOM URL：[黄色:〇〇〇[/黄色]",
        ],
        "yellow_fields": ["冒頭の確認差し替え欄", "ZOOM URL：〇〇〇"],
        "social_style_variants": {
            "Driver / Driving": "件名から署名まで含む全文メール案を入れます。",
            "Analytical": "件名から署名まで含む全文メール案を入れます。",
            "Amiable": "件名から署名まで含む全文メール案を入れます。",
            "Expressive": "件名から署名まで含む全文メール案を入れます。",
        },
        "deal_feedback": {
            "総合概要": {
                "overallSummary": "顧客は費用と作業時間を現実的に確認しながら検討している。商談後メールでは、判断材料を整理し、次回接点で不明点を回収する導線を作る。",
                "meetingOutcome": "検討中",
                "benchmarkGapAnalysis": {
                    "keyDifference": "価格不安を、初期費用そのもの・回収可能性・作業時間に切り分ける余地がある。",
                    "missingActions": [
                        {"phase": "不安の特定", "missing": "不安の種類の切り分け", "impact": "価格だけで判断されやすい"},
                    ],
                },
            },
            "顧客インサイト": {
                "hiddenNeeds": ["無理なく続けられるか確認したい"],
                "unspokenIssues": ["家族に説明できる材料が不足している可能性"],
                "competingServices": ["明確な競合名は未確認"],
                "decisionCriteria": ["費用", "作業時間", "サポート範囲"],
                "buyingTemperature": "中: 判断材料を求めているが即決ではない",
            },
            "認知バイアス": [
                {
                    "name": "損失回避",
                    "description": "初期費用の失敗リスクを強く見ている可能性がある。",
                    "counterMeasure": "費用、作業量、サポート範囲を分けて確認する。",
                },
            ],
            "期待値のズレ": [
                {
                    "topic": "回収時期",
                    "gap": "早期回収への期待と現実の作業量にズレが出る可能性がある。",
                    "solution": "作業時間と回収イメージを分けて説明する。",
                },
            ],
            "良かった点": [
                {
                    "item": "不安の受け止め",
                    "detail": "価格への反応を否定せず、判断材料へ戻す方向性がある。",
                },
            ],
            "改善ポイント": [
                {
                    "item": "次回質問の明確化",
                    "detail": "費用・作業時間・家族説明のどれが最重要かを確認する。",
                },
            ],
            "AIコーチングカード": [
                {
                    "theme": "費用不安の切り分け",
                    "scene": "顧客が金額を気にした場面",
                    "insight": "価格だけで判断するとプラン選定がずれやすい。",
                    "issue": "不安の種類が未分解。",
                    "strategy": "初期費用、回収可能性、作業時間を分けて聞く。",
                    "script": "金額はもちろん大事な判断材料です。ただ、金額だけで見ると合うプランを間違えやすいので、作業時間と目的を先に確認させてください。",
                    "outcome": "顧客の本音を把握し、提案の焦点を合わせやすくなる。",
                },
            ],
            "再現する勝ち筋": [
                {
                    "name": "価格不安を判断条件へ変換する",
                    "useCase": "商談初期に金額を聞かれた時",
                    "customerSignals": ["高いですか？", "回収できますか？"],
                    "topSalesMove": ["受け止める", "不安の種類を切り分ける", "目的と作業時間へ戻す"],
                    "copyTalk": "そこは皆さん一番気にされます。まず、初期費用そのものが不安なのか、回収できるかが不安なのかを分けて確認させてください。",
                    "whyItWorks": "価格だけの比較を避け、顧客の判断条件を具体化できる。",
                    "ngExamples": ["大丈夫ですとだけ返す", "価格だけ即答して終わる"],
                    "practice": "価格質問への返答を3パターンの短い返答案として用意する。",
                },
            ],
            "商談フェーズ": {
                "currentPhase": "不安の特定",
                "currentGoal": "判断材料を整理し、次回質問につなげる",
                "mustHearBeforeProposal": ["作業時間", "予算感", "意思決定者"],
            },
            "次の一手": {
                "nextBestAction": "送信後、資料確認の有無と不明点を確認する。",
                "hearingQuestions": ["一番気になるのは費用面ですか、作業面ですか？"],
                "recommendedAnswer": "金額だけで判断するとズレやすいので、作業時間と目的に合うかを一緒に確認させてください。",
            },
            "ベンチマーク営業再現": {
                "script": "そこは皆さん一番気にされます。まず状況を確認した上で、合うかどうかを一緒に見ていきましょう。",
                "delivery": "急かさず、短く区切って伝える。",
            },
        },
        "final_check": {"顧客送付用本文": "確認済み", "黄色箇所": "確認済み"},
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, help="JSON input file")
    parser.add_argument("--output", type=Path, required=True, help="Output .docx path")
    parser.add_argument("--self-test", action="store_true", help="Use built-in sample data")
    args = parser.parse_args()

    if args.self_test:
        data = sample_data()
    elif args.input:
        data = json.loads(args.input.read_text(encoding="utf-8"))
    else:
        parser.error("Use --input or --self-test")

    build_document(data, args.output)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
