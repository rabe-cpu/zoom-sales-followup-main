#!/usr/bin/env python3
"""Register or propose updates to knowledge/sales_persons/{name}.md and {name}.docx.

This script always emits BOTH the Markdown and the Word (.docx) version, so the
human-review docx stays in sync with the machine-readable Markdown that
CLAUDE.md imports.

Modes:
- new:   Create {name}.md AND {name}.docx, then append @{name}.md to _index.md.
         Fails if the canonical .md already exists.
- draft: Write {name}.draft.md AND {name}.draft.docx for human diff review.
         Never modifies the canonical files or _index.md.
- auto:  Pick new if the canonical .md does not exist, otherwise draft.
- regen: Re-render the .docx from an existing .md without rewriting the .md.
         Useful for backfilling docx for files that were created before the
         dual-output rule. Requires --md-path.

Input for new/draft/auto: --input person.json (see SKILL.md for schema).
Input for regen: --md-path existing.md (no JSON needed).
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


REQUIRED_KEYS = ["name"]
INDEX_FILENAME = "_index.md"
INDEX_MARKER = "## 登録済み"
BLUE = "D9EAF7"
GREEN = "E2F0D9"


def render_md(data: dict[str, Any]) -> str:
    today = dt.date.today().isoformat()
    name = data["name"]
    display_name = data.get("display_name", name)
    company = data.get("company", "")
    sources = data.get("source_transcripts", []) or []
    features = data.get("features", []) or []
    first_person = data.get("first_person", "")
    customer_address = data.get("customer_address", "")
    email_phrases = data.get("email_phrases", []) or []
    closing = data.get("closing", "")
    slack = data.get("slack_expressions", {}) or {}
    prohibitions = data.get("prohibitions", []) or []

    lines: list[str] = []
    lines.append("---")
    lines.append(f"name: {name}")
    lines.append(f"display_name: {display_name}")
    if company:
        lines.append(f"company: {company}")
    lines.append(f"created: {today}")
    lines.append(f"updated: {today}")
    lines.append("source_transcripts:")
    if sources:
        for s in sources:
            lines.append(f"  - {s}")
    else:
        lines.append("  []")
    lines.append("---")
    lines.append("")
    lines.append(f"# {display_name}さん")
    lines.append("")
    lines.append("## 商談内で見えた特徴")
    lines.append("")
    if features:
        for f in features:
            lines.append(f"- `{f}`")
    else:
        lines.append("- （抽出根拠なし）")
    lines.append("")
    lines.append("## 一人称・距離感")
    lines.append("")
    lines.append(f"- 一人称: {first_person or '（未確認）'}")
    lines.append(f"- 顧客への呼び方: {customer_address or '（未確認）'}")
    lines.append("")
    lines.append("## メール反映の定型")
    lines.append("")
    if email_phrases:
        for p in email_phrases:
            lines.append(f"- `{p}`")
    else:
        lines.append("- （未抽出）")
    if closing:
        lines.append(f"- 締め: `{closing}`")
    lines.append("")
    lines.append("## 余白表現の出現")
    lines.append("")
    for key in ["かなと", "かと", "できればと", "見ていただけると"]:
        value = slack.get(key, "（未観測）")
        lines.append(f"- {key}: {value}")
    lines.append("")
    lines.append("## 禁止候補")
    lines.append("")
    if prohibitions:
        for p in prohibitions:
            lines.append(f"- `{p}`")
    else:
        lines.append("- （該当なし）")
    lines.append("")
    return "\n".join(lines)


def parse_md(text: str) -> dict[str, Any]:
    """Parse a knowledge/sales_persons/*.md back into the structured form."""
    lines = text.splitlines()
    frontmatter: dict[str, Any] = {}
    sources: list[str] = []
    body_start = 0

    if lines and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            stripped = lines[index].rstrip()
            if stripped == "---":
                body_start = index + 1
                break
            if stripped.startswith("source_transcripts:"):
                rest = stripped.split(":", 1)[1].strip()
                if rest and rest != "[]":
                    sources.append(rest)
                continue
            if stripped.startswith("  - "):
                sources.append(stripped[4:].strip())
                continue
            if ":" in stripped:
                key, value = stripped.split(":", 1)
                frontmatter[key.strip()] = value.strip()

    sections: dict[str, list[str]] = {}
    title = frontmatter.get("display_name", frontmatter.get("name", ""))
    current = None
    bucket: list[str] = []
    for line in lines[body_start:]:
        if line.startswith("# "):
            title = line[2:].strip().rstrip("さん")
        elif line.startswith("## "):
            if current is not None:
                sections[current] = bucket
            current = line[3:].strip()
            bucket = []
        else:
            bucket.append(line)
    if current is not None:
        sections[current] = bucket

    return {
        "title": title,
        "frontmatter": frontmatter,
        "sources": sources,
        "sections": sections,
    }


def set_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Yu Gothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    style.font.size = Pt(10.5)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_kv_table(document: Document, headers: list[str], rows: list[list[str]], fill: str) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        shade_cell(cell, fill)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)


def add_bullets(document: Document, items: list[str]) -> None:
    for raw in items:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("- "):
            text = line[2:].strip().replace("`", "")
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(line.replace("`", ""))


def render_docx_from_md(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    parsed = parse_md(text)

    document = Document()
    set_font(document)

    title_text = f"{parsed['title']}さん 口調ナレッジ"
    heading = document.add_heading(title_text, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("基本情報", level=1)
    fm = parsed["frontmatter"]
    rows = [
        ["name", fm.get("name", "")],
        ["display_name", fm.get("display_name", "")],
        ["company", fm.get("company", "")],
        ["created", fm.get("created", "")],
        ["updated", fm.get("updated", "")],
        ["source_transcripts", "\n".join(parsed["sources"]) if parsed["sources"] else "（なし）"],
    ]
    add_kv_table(document, ["項目", "内容"], rows, BLUE)

    for heading_text, body_lines in parsed["sections"].items():
        document.add_heading(heading_text, level=1)
        add_bullets(document, body_lines)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def update_index(index_path: Path, name: str) -> bool:
    if not index_path.exists():
        index_path.write_text(
            "# Sales Persons Index\n\n"
            "営業担当ごとの口調・余白表現・禁止候補は、このフォルダの個別ファイルに分けて管理する。\n\n"
            f"{INDEX_MARKER}\n\n@{name}.md\n",
            encoding="utf-8",
        )
        return True

    text = index_path.read_text(encoding="utf-8")
    line = f"@{name}.md"
    if line in text:
        return False

    if INDEX_MARKER in text:
        new_text = text.rstrip() + f"\n{line}\n"
    else:
        new_text = text.rstrip() + f"\n\n{INDEX_MARKER}\n\n{line}\n"
    index_path.write_text(new_text, encoding="utf-8")
    return True


def validate(data: dict[str, Any]) -> None:
    for key in REQUIRED_KEYS:
        if key not in data or not str(data[key]).strip():
            raise SystemExit(f"input missing required key: {key}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", type=Path, help="JSON describing the salesperson (required for new/draft/auto)")
    parser.add_argument(
        "--knowledge-dir",
        type=Path,
        required=True,
        help="Path to knowledge/sales_persons/",
    )
    parser.add_argument("--mode", choices=["new", "draft", "auto", "regen"], default="auto")
    parser.add_argument("--md-path", type=Path, help="Existing .md to regenerate .docx from (regen mode only)")
    return parser.parse_args()


def emit_pair(data: dict[str, Any], md_path: Path, docx_path: Path) -> None:
    md_text = render_md(data)
    md_path.write_text(md_text, encoding="utf-8")
    render_docx_from_md(md_path, docx_path)


def main() -> int:
    args = parse_args()
    knowledge_dir: Path = args.knowledge_dir
    knowledge_dir.mkdir(parents=True, exist_ok=True)

    if args.mode == "regen":
        if not args.md_path:
            raise SystemExit("regen mode requires --md-path")
        md_path: Path = args.md_path
        if not md_path.exists():
            raise SystemExit(f"{md_path} does not exist")
        docx_path = md_path.with_suffix(".docx")
        render_docx_from_md(md_path, docx_path)
        print(json.dumps({"mode": "regen", "md": str(md_path), "docx": str(docx_path)}, ensure_ascii=False))
        return 0

    if not args.input:
        raise SystemExit("new/draft/auto mode requires --input")
    data = json.loads(args.input.read_text(encoding="utf-8"))
    validate(data)

    name = data["name"].strip()
    canonical_md = knowledge_dir / f"{name}.md"
    canonical_docx = knowledge_dir / f"{name}.docx"
    draft_md = knowledge_dir / f"{name}.draft.md"
    draft_docx = knowledge_dir / f"{name}.draft.docx"

    mode = args.mode
    if mode == "auto":
        mode = "draft" if canonical_md.exists() else "new"

    if mode == "new":
        if canonical_md.exists():
            raise SystemExit(
                f"{canonical_md} already exists. Use --mode draft to propose changes or --mode regen to refresh docx."
            )
        emit_pair(data, canonical_md, canonical_docx)
        index_changed = update_index(knowledge_dir / INDEX_FILENAME, name)
        print(json.dumps(
            {
                "mode": "new",
                "md": str(canonical_md),
                "docx": str(canonical_docx),
                "index_updated": index_changed,
            },
            ensure_ascii=False,
        ))
        return 0

    if mode == "draft":
        emit_pair(data, draft_md, draft_docx)
        print(json.dumps(
            {
                "mode": "draft",
                "md": str(draft_md),
                "docx": str(draft_docx),
                "canonical_md": str(canonical_md),
                "canonical_exists": canonical_md.exists(),
            },
            ensure_ascii=False,
        ))
        return 0

    raise SystemExit(f"unknown mode: {mode}")


if __name__ == "__main__":
    raise SystemExit(main())
